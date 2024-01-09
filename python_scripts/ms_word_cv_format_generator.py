from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL #WD_ALIGN_HORIZONTAL

import json,yaml,time,os,glob
import streamlit as st

class FormatMSword():

    def __init__(self) -> None:
        self.document=Document()
 
    def read_yaml(self,path):
        with open(path, 'r') as yaml_file:
            return yaml.safe_load(yaml_file)

    def get_full_path(self,section):
        # Define the relative path to the resume section files
        relative_path = "./section_yaml_files"
        # # Create the full path to the output folder
        # output_folder_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), relative_path)
        # section_file_path_yaml = f'{output_folder_path}/resume_{section}.yaml'
        section_file_path_yaml = f'{relative_path}/resume_{section}.yaml'
        return section_file_path_yaml

    def heading(self,text, level, bold=True, color=(0, 0, 0),align=WD_ALIGN_PARAGRAPH.LEFT):
        '''Re-writing add_heading function of Document() for 
            our requirements.'''
        heading_added = self.document.add_heading(text, level=level)
        heading_added.alignment =align
        run = heading_added.runs[0]
        run.bold = bold
        run.font.color.rgb = RGBColor(*color)
        return heading_added


# Caching resource in order to reuse the word file generated at the end;
# so that there is no change in job requirement on re-running.
# @st.cache_resource
def create_ms_word_doc():
    # Create a new Word document instance
    doc=FormatMSword()

    # Add a title and contact
    doc.heading('YOUR NAME', level=0,bold=True, color=(255, 0, 0))
    doc.heading('0422332233, email@gmail.com', level=9,bold=False,align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Specify the yaml file path of all sections
    file_path_exp = doc.get_full_path('experience')
    file_path_skill = doc.get_full_path('skill')
    file_path_education = doc.get_full_path('education')
    file_path_projects = doc.get_full_path('projects')
    file_path_awards = doc.get_full_path('awards')

    # Read the JSON file back into a dictionary
    exp_details = doc.read_yaml(file_path_exp)
    skill_details = doc.read_yaml(file_path_skill)
    education_details = doc.read_yaml(file_path_education)
    projects_details = doc.read_yaml(file_path_projects)
    awards_details = doc.read_yaml(file_path_awards)

    # level 1 heading color
    level_1_heading_color=(255, 0, 0)

    # Add EXPERIENCE section
    doc.heading('EXPERIENCE', level=1, bold=True, color=level_1_heading_color)
    # Iterate through experience details
    for exp in exp_details[list(exp_details.keys())[0]]:
        for k,v in exp.items():

            if k.lower().strip() == 'employer':
                key = doc.heading((k.upper()), level=3,color=(0, 32, 96))

                if isinstance(v,list):
                    for line in v:
                        doc.document.add_paragraph(line,style='List Bullet')
                else:
                    run= key.add_run(f'     {v}')
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                key = doc.heading((k.upper()), level=3, color=(0, 32, 96))

                if isinstance(v,list):
                    for line in v:
                        doc.document.add_paragraph(line,style='List Bullet')
                else:
                    run= key.add_run(f'     {v}')
                    run.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
        doc.document.add_paragraph()   # to add empty line
    # doc.document.add_page_break() # Add Page break after experience


    # Skill section
    skills = doc.heading('SKILLS', level=1,bold=True, color=level_1_heading_color)
    skills.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a table with two columns
    table = doc.document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    table.columns.autofit = False

    # Set column widths
    table.columns[0].width = Pt(50)
    table.columns[1].width = Pt(400)

    # Add header row
    header_row = table.rows[0].cells
    header_row[0].text = 'Key Skill'
    header_row[1].text = 'Skill Description'

    # Make the table header bold
    for cell in header_row:
        cell.paragraphs[0].runs[0].bold = True

    skills_data=skill_details[list(skill_details.keys())[0]]
    # Add data to table
    for skill in skills_data:
        row_cells = table.add_row().cells
        skill_key=list(skill.keys())
        # Make the skill name bold
        skill_name_cell = row_cells[0].paragraphs[0].add_run(skill[skill_key[0]])
        skill_name_cell.bold = True
        row_cells[1].text = skill[skill_key[1]]
    # doc.document.add_paragraph()   # to add empty line after the section


    # Projects section
    projects = doc.heading('PROJECTS', level=1,bold=True, color=level_1_heading_color)
    projects.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for _,v in projects_details[list(projects_details.keys())[0]].items():
        doc.document.add_paragraph(v,style='List Bullet')
    # doc.document.add_paragraph()   # to add empty line


    # Education section
    edu = doc.heading('EDUCATION', level=1,bold=True, color=level_1_heading_color)
    edu.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for edu in education_details[list(education_details.keys())[0]]:
        for k,v in edu.items():
            key = doc.heading((k.upper()), level=3,color=(0, 32, 96))
            key.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(v,list):
                for line in v:
                    doc.document.add_paragraph(v,style='List Bullet')
            else:
                run= key.add_run(f'     {v}')
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
    # doc.document.add_paragraph()   # to add empty line


    # Awards section
    awards = doc.heading('AWARDS', level=1,bold=True, color=level_1_heading_color)
    awards.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for _,v in awards_details[list(awards_details.keys())[0]].items():
        doc.document.add_paragraph(v,style='List Bullet')
    # doc.document.add_paragraph()   # to add empty line


    # Add footer
    footer = doc.document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    run.add_text('Resume - Your Name')
    run.font.size = Pt(8)


    
    # Note down relative path. './' is current workdir and '../' is to go one level up to the parent folder.
    relative_output_path = "./generated_output_files"

    # remove any previously generated file from the folder
    files_to_delete = glob.glob(f"{relative_output_path}/resume_ideation_msword_*.docx")
    if len(files_to_delete)>0:
        for file_path in files_to_delete:
            os.remove(file_path)
        print("Previously existing files have been deleted.")
    else:
        print("No previously existing files found.")

    # generate doc at path
    output_file_path=f'{relative_output_path}/resume_ideation_msword_{time.time()}.docx'
    doc.document.save(output_file_path)

    return output_file_path


if __name__ == '__main__':
    output_file_path=create_ms_word_doc()
    print(f'\nSuccessfully created resume ideation word file at {output_file_path}')